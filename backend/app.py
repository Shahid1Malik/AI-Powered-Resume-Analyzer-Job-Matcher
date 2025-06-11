# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from sklearn.feature_extraction.text import TfidfVectorizer
# import PyPDF2
# import docx

# app = Flask(__name__)
# CORS(app)

# def extract_text(file):
#     if file.filename.endswith('.pdf'):
#         reader = PyPDF2.PdfReader(file)
#         return " ".join(page.extract_text() for page in reader.pages if page.extract_text())
#     elif file.filename.endswith('.docx'):
#         doc = docx.Document(file)
#         return " ".join([para.text for para in doc.paragraphs])
#     else:
#         return ""

# @app.route('/match', methods=['POST'])
# def match_resume():
#     resume_file = request.files.get('resume')
#     job_description = request.form.get('job')
#     resume_text = extract_text(resume_file)
#     print("Extracted Resume Text:", resume_text)
#     print("Job Description:", job_description)


#     if not resume_file or not job_description:
#         return jsonify({'error': 'Missing file or job description'}), 400

#     resume_text = extract_text(resume_file)

#     vectorizer = TfidfVectorizer().fit([resume_text, job_description])
#     vectors = vectorizer.transform([resume_text, job_description])
#     score = (vectors[0] @ vectors[1].T).toarray()[0][0]

#     return jsonify({'match_score': round(score * 100, 2)})



# if __name__ == '__main__':
#     app.run(debug=True)


from flask import Flask, request, jsonify
from flask_cors import CORS
from sklearn.feature_extraction.text import TfidfVectorizer
import PyPDF2
import docx
import os
import requests
from dotenv import load_dotenv
load_dotenv()
app = Flask(__name__)
CORS(app)

def extract_text(file):
    if file.filename.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return " ".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif file.filename.endswith('.docx'):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])
    else:
        return ""

@app.route('/match', methods=['POST'])
def match_resume():
    resume_file = request.files.get('resume')
    job_description = request.form.get('job')

    if not resume_file or not job_description:
        return jsonify({'error': 'Missing file or job description'}), 400

    resume_text = extract_text(resume_file)
    print("Extracted Resume Text:", resume_text[:200])  # limit print for readability
    print("Job Description:", job_description[:200])

    vectorizer = TfidfVectorizer().fit([resume_text, job_description])
    vectors = vectorizer.transform([resume_text, job_description])
    score = (vectors[0] @ vectors[1].T).toarray()[0][0]

    return jsonify({'match_score': round(score * 100, 2)})

# ✅ NEW: Fetch jobs from Proxycurl or dummy data fallback
@app.route('/jobs', methods=['GET'])
def get_jobs():
    api_key = os.getenv('PROXYCURL_API_KEY')
    if not api_key:
        # Fallback dummy data
        return jsonify([
            {
                "job_title": "Frontend Developer (React)",
                "location": "Remote",
                "job_url": "https://example.com/job/react-dev"
            },
            {
                "job_title": "Machine Learning Engineer",
                "location": "New York, NY",
                "job_url": "https://example.com/job/ml-engineer"
            }
        ])

    headers = {'Authorization': f'Bearer {api_key}'}
    params = {'search_id': '1035'}  # Example: Apple Inc LinkedIn ID
    response = requests.get('https://nubela.co/proxycurl/api/v2/linkedin/company/job', headers=headers, params=params)

    if response.status_code != 200:
        return jsonify({'error': 'Failed to fetch jobs'}), response.status_code

    data = response.json()
    return jsonify(data.get('job', []))

if __name__ == '__main__':
    app.run(debug=True)
